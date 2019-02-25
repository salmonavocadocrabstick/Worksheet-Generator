
# Tools for building a sentence.
import sentence_obj_builder as s_obj_builder			 	# Creates elements
import sentence_template as s_obj_template				# Templates for different sentence types
import sentence_maker as s_make						# Generates the sentence.

from docx import Document
from docx.shared import Pt #for font size
from docx.enum.style import WD_STYLE_TYPE
from datetime import date


def _FIB(FIB_obj, count):
	FIB_sentence = s_make.generate_sentence_by_type(FIB_obj)
	print(FIB_sentence.get_full_sentence())	#Console will show what the sentence is.
	p = document.add_paragraph(f"{count}. {FIB_sentence.get_part_a()}", style = "Normal")
	p.add_run("______________").font.name = "Comic Sans"
	p.add_run(f"{FIB_sentence.get_part_b()}")

def Fill_In_The_Blanks_Verb(s_obj, count):
	FIBV_obj = s_obj_template.FIBV_sentence_obj(s_obj)
	_FIB(FIBV_obj, count)

def Fill_In_The_Blanks_Noun(s_obj, count):
	FIBV_obj = s_obj_template.FIBN_sentence_obj(s_obj)
	_FIB(FIBV_obj, count)

def Answering_Question(s_obj, count):
	AQ_obj = s_obj_template.Q_sentence_obj(s_obj)
	AQ_sentence = s_make.generate_sentence_by_type(AQ_obj)
	print(AQ_sentence.get_full_sentence())

	document.add_paragraph(f"{count}. {AQ_sentence.get_full_sentence()}", style = "Normal")
	document.add_paragraph("________________________________________________")

def Finding_Mistakes(s_obj, count):
	pass #Under construction.

def Rearrange_Sentence(s_obj, count):
	RS_obj = s_obj_template.FS_sentence_obj(s_obj)
	RS_sentence = s_make.generate_sentence_by_type(RS_obj)
	print(RS_sentence.get_full_sentence())

	## Getting a list of shuffled words, just pop them one by one in the for loop!.
	list_of_words = RS_sentence.get_jumboed_sentence()
	slash = " / "
	slash = slash.join(list_of_words)

	document.add_paragraph(f"{count}. {slash}", style = "Normal")
	document.add_paragraph("________________________________________________")



if __name__ == "__main__":
	#Variables

	# Change document title here:
	doc_title = "Verbs"

	# Choose Instruction from here:
	instruction = {
						"i0":	"Fill in the blanks with the correct verb.", #0
						"i1":	"Fill in the blanks with the correct noun.",
						"i2":	"Rearrange the sentences into the correct order.",
						"i3":	"Answer the following questions in full sentences."

					}


	file_name = {

						"i0":	"FIBV_", #0
						"i1":	"FIBN_",
						"i2":	"RS_",
						"i3":	"AQ_"

				}


	part_a_instru = instruction["i0"]
	part_b_instru = instruction["i3"]

	#Decide number of questions :
	part_a_num_questions = 15
	part_b_num_questions = 15


	document = Document()


	styles = document.styles

	# Title:
	heading_style = styles.add_style('Heading St', WD_STYLE_TYPE.PARAGRAPH)
	heading_style.base_style = styles["Heading 1"]
	heading_style.font.name = "Janda Manatee Bubble"
	heading_style.font.size = Pt(30)

	document.add_paragraph(doc_title, style = "Heading St")

	#Document Styling
	style = document.styles["Normal"]
	font = style.font
	font.name = "Architect"
	font.size = Pt(18)


	document.add_heading(part_a_instru, level=2)
	for x in range(1, part_a_num_questions+1):
		s_obj = s_obj_builder.get_s_obj()

		### Run the function here ###
		Answering_Question(s_obj, x)
		#Rearrange_Sentence(s_obj, x)
		#Fill_In_The_Blanks_Noun(s_obj, x)
		#Fill_In_The_Blanks_Verb(s_obj, x)

	document.add_page_break()
	document.add_heading(part_b_instru, level=2)
	for y in range(1, part_b_num_questions+1):
		s_obj = s_obj_builder.get_s_obj()

		### Run the function here ###
			#Answering_Question(s_obj, y)
			#Rearrange_Sentence(s_obj, y)
		Fill_In_The_Blanks_Verb(s_obj, y)
			#Fill_In_The_Blanks_Verb(s_obj, y)
	
	file_save = file_name["i0"] + file_name["i1"]
	document.save(f"{file_save}{date.today()}.docx")
	print("----------------------------------------------------------------------------")
	print(f"Worksheet {file_save}{date.today()}.docx is generated in designated folder.")