import lambo_obj as obj
import lambo_verb as vrb
import lambo_sub as subj
import present_keywords as pre


from docx import Document
from docx.shared import Pt #for font size
from datetime import date

#Variables
doc_title = "Verbs"
instruction = [
					"Fill in the blanks with the correct verb.", #0
					"Fill in the blanks with the correct noun."

				]
doc_instru = instruction[0]
num_questions = 20


document = Document()
document.add_heading(doc_title, 0)

#Document Styling
style = document.styles["Normal"]
font = style.font
font.name = "KG Neatly Printed Spaced"
font.size = Pt(18)



document.add_heading(doc_instru, level=1)


for x in range(1, num_questions+1):
	verb_list = vrb.verb()
	verb = verb_list[1]
	#print(verb)
	adj_noun = obj.object(verb_list[0])
	#print(adj_noun)
	
	subject = subj.subject()
	keyword = pre.present_tense_keyword()

	#Fill in the blanks
	p = document.add_paragraph(f"{x}. {subject} ", style = "Normal")
	p.add_run("______________").font.name = "Comic Sans"
	p.add_run(f"({verb}) {adj_noun} {keyword}.")


	#document.add_paragraph(this_question, style='Normal')

document.add_page_break()
document.save(f"{doc_title}_{date.today()}.docx")